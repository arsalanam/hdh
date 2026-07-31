"""Generate a Word feature guide for the Family Medicine Synthetic Dataset project.

Sources: README.md and ARCHITECTURE.md.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette --------------------------------------------------------------
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TEAL = RGBColor(0x16, 0x7D, 0x7F)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = "EAF1F4"

doc = Document()

# ---- base styles ----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)

for name, size, color in [
    ("Heading 1", 18, NAVY),
    ("Heading 2", 14, TEAL),
    ("Heading 3", 12, NAVY),
]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10.5)
        shade_cell(hdr[i], "1F3A5F")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for r in table.rows:
                r.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # light shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), LIGHT)
    pPr.append(shd)
    return p


def add_bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def feature(title, what, value):
    h = doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    p.add_run("What it does:  ").bold = True
    p.add_run(what)
    p = doc.add_paragraph()
    p.add_run("Why it matters:  ").bold = True
    p.add_run(value)


# ===========================================================================
# COVER
# ===========================================================================
for _ in range(3):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Family Medicine Synthetic Dataset")
r.font.size = Pt(30)
r.font.bold = True
r.font.color.rgb = NAVY

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Feature Guide")
r.font.size = Pt(20)
r.font.color.rgb = TEAL

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("10,000 patients  ·  165,000+ visits  ·  777,000+ lab results")
r.font.size = Pt(13)
r.font.color.rgb = GREY

for _ in range(8):
    doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("A medically realistic synthetic OPD dataset for testing agentic AI care programs")
r.italic = True
r.font.color.rgb = GREY

dt = doc.add_paragraph()
dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dt.add_run("Stack: Python · Faker · SQLAlchemy · SQLite · Alembic")
r.font.size = Pt(10)
r.font.color.rgb = GREY

doc.add_page_break()

# ===========================================================================
# TABLE OF CONTENTS (static)
# ===========================================================================
doc.add_heading("Contents", level=1)
toc = [
    "1.  Overview",
    "2.  Key Features at a Glance",
    "3.  Getting Started",
    "4.  Data Model",
    "5.  The Disease Engine",
    "6.  Command-Line Interface",
    "7.  Export Formats",
    "8.  Extensible Schema Architecture",
    "9.  Dataset Statistics",
    "10. Roadmap & Possible Extensions",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(3)
doc.add_page_break()

# ===========================================================================
# 1. OVERVIEW
# ===========================================================================
doc.add_heading("1.  Overview", level=1)
doc.add_paragraph(
    "The Family Medicine Synthetic Dataset is a generator and pre-built database of "
    "medically realistic outpatient (OPD) records. It produces 10,000 synthetic patients "
    "with roughly four years of visit history — 165,000+ visits, 165,000+ diagnoses, "
    "170,000+ prescriptions, and 777,000+ lab results — all driven by age, sex, and "
    "seasonal disease probabilities."
)
doc.add_paragraph(
    "The dataset is purpose-built as a safe, no-PHI sandbox for developing and testing "
    "agentic AI care programs, clinical decision tools, FHIR pipelines, and analytics — "
    "without touching real patient data."
)
doc.add_heading("Who it's for", level=3)
add_bullets([
    "AI/ML engineers building or evaluating agentic care programs.",
    "Healthcare integration developers testing FHIR R4 pipelines.",
    "Data scientists prototyping risk-stratification or care-gap models.",
    "Anyone needing realistic clinical data without privacy or compliance overhead.",
])

# ===========================================================================
# 2. KEY FEATURES
# ===========================================================================
doc.add_heading("2.  Key Features at a Glance", level=1)
add_table(
    ["Feature", "Description"],
    [
        ["Realistic generation", "Age/sex/season-weighted disease probability engine with comorbidity seeding."],
        ["30 conditions", "Coverage spanning pediatric, adolescent, adult, and senior care."],
        ["Full clinical detail", "Vitals, ICD-10 diagnoses, formulary-accurate prescriptions, and LOINC-coded labs."],
        ["Three export formats", "Per-patient JSON, FHIR R4 Bundles, and LLM-ready plain-text charts."],
        ["Maintenance CLI", "Generate, inspect, export, inject disease spikes, and advance time."],
        ["Pre-built database", "Ships with a ~87 MB SQLite database of 10,000 patients."],
        ["Extensible schema", "JSON-driven, modular schema registry with Alembic migration support."],
    ],
    widths=[2.0, 4.5],
)

# ===========================================================================
# 3. GETTING STARTED
# ===========================================================================
doc.add_heading("3.  Getting Started", level=1)
doc.add_paragraph("Install dependencies and generate a dataset in a few commands:")
add_code(
    "pip install -r requirements.txt\n\n"
    "# Generate 10,000 patients (4 years of history)\n"
    "python cli.py generate --patients 10000 --years 4\n\n"
    "# View statistics\n"
    "python cli.py stats\n\n"
    "# Export to JSON / FHIR R4 / Plain text\n"
    "python cli.py export --format all --limit 500 --output-dir exports/\n\n"
    "# Show one patient chart\n"
    "python cli.py show --mrn MRN12345678"
)
doc.add_paragraph(
    "A pre-generated database (family_medicine.db) is included, so you can run stats, "
    "show, and export commands immediately without regenerating."
)

doc.add_heading("Project Files", level=3)
add_table(
    ["File", "Purpose"],
    [
        ["models.py", "SQLAlchemy ORM — Patient, Visit, Vital, Diagnosis, Prescription, LabResult."],
        ["disease_engine.py", "Age/sex/season probability engine, ICD-10 codes, medication formularies."],
        ["generators.py", "Patient and visit-history generators."],
        ["exporters.py", "JSON, FHIR R4 Bundle, and plain-text clinical-note exporters."],
        ["cli.py", "Maintenance command-line interface."],
        ["family_medicine.db", "Pre-generated SQLite database (10,000 patients)."],
    ],
    widths=[1.8, 4.7],
)

# ===========================================================================
# 4. DATA MODEL
# ===========================================================================
doc.add_heading("4.  Data Model", level=1)
doc.add_paragraph(
    "The schema is defined in models.py as SQLAlchemy ORM classes. Each patient anchors "
    "a tree of visits, and each visit carries its own vitals, diagnoses, prescriptions, "
    "and labs."
)
add_table(
    ["Model", "Table", "Key Columns"],
    [
        ["Patient", "patients", "mrn, dob, sex, race, insurance, allergies, fam_hx_*, smoker, bmi_baseline"],
        ["ChronicCondition", "chronic_conditions", "patient_id, icd10_code, onset_date, controlled"],
        ["Visit", "visits", "patient_id, visit_date, visit_type, chief_complaint, provider_name, follow_up_days"],
        ["Vital", "vitals", "visit_id, bp_systolic, bp_diastolic, hr, rr, temp_f, spo2, weight_kg, bmi, pain_scale"],
        ["Diagnosis", "diagnoses", "visit_id, icd10_code, description, is_primary"],
        ["Prescription", "prescriptions", "visit_id, drug_name, drug_class, dose, frequency, duration_days, refills, is_new"],
        ["LabResult", "lab_results", "visit_id, test_name, value, unit, ref_low, ref_high, status, loinc_code"],
    ],
    widths=[1.4, 1.4, 3.7],
)

doc.add_heading("Enums", level=3)
add_bullets([
    "Sex: M / F",
    "VisitType: acute / follow_up / preventive / urgent",
    "LabStatus: normal / high / low / critical",
])

doc.add_heading("Key Relationships", level=3)
add_code(
    "Patient ──< Visit ──< Diagnosis\n"
    "                 ──< Prescription\n"
    "                 ──< LabResult\n"
    "                 ──1 Vital\n"
    "Patient ──< ChronicCondition"
)

# ===========================================================================
# 5. DISEASE ENGINE
# ===========================================================================
doc.add_heading("5.  The Disease Engine", level=1)
doc.add_paragraph(
    "The disease engine (disease_engine.py) is the heart of the realism. Each condition is "
    "modeled as a ConditionProfile dataclass that defines exactly how a visit for that "
    "condition should look."
)
doc.add_heading("Each ConditionProfile defines", level=3)
add_bullets([
    "icd10_code, description, chief_complaint, and visit_type.",
    "Vital deltas from baseline (mean, sd) for BP, HR, RR, temperature, SpO2, and pain.",
    "labs — which lab panels to order (LabSpec).",
    "rx_options — condition-appropriate formulary entries (RxSpec).",
    "follow_up_days — a clinical-guideline follow-up interval.",
    "seasonal_weights — month-to-multiplier seasonal weighting.",
])

doc.add_heading("30 Conditions by Age Group", level=3)
add_table(
    ["Age Group", "Conditions"],
    [
        ["Infant 0–2", "Well-child, otitis media, RSV, febrile illness, rash/eczema, conjunctivitis, URI"],
        ["Child 3–12", "Well-child, otitis media, strep throat, URI, febrile illness, rash, conjunctivitis, sports injury"],
        ["Teen 13–17", "Sports physical, URI, sports injury, acne, anxiety, strep, mono"],
        ["Young Adult 18–35", "Annual physical, influenza, URI, UTI, anxiety, low back pain, laceration, contraception, GERD"],
        ["Adult 36–50", "Annual physical, HTN, hyperlipidemia, T2DM, URI, influenza, GERD, anxiety, back pain, obesity"],
        ["Middle-aged 51–65", "Annual physical, HTN, T2DM, hyperlipidemia, osteoarthritis, GERD, URI, COPD, depression, hypothyroidism"],
        ["Senior 65+", "Annual wellness, HTN, T2DM, hyperlipidemia, osteoarthritis, COPD, falls, polypharmacy review, depression, hypothyroidism, influenza"],
    ],
    widths=[1.6, 4.9],
)

doc.add_heading("Seasonal Multipliers", level=3)
doc.add_paragraph(
    "Disease incidence is weighted by month so the dataset shows realistic seasonality."
)
add_code(
    "FLU_SEASON  = {Jan:2.5, Feb:2.0, ..., Dec:2.5}   # peaks winter\n"
    "RSV_SEASON  = {Jan:2.0, Feb:1.5, ..., Dec:2.5}   # peaks late fall/winter\n"
    "SUMMER_PEAK = {Jun:1.5, Jul:1.5, Aug:1.5, ...}   # UTI, sports injuries, lacerations"
)

doc.add_heading("Comorbidity Seeding", level=3)
doc.add_paragraph(
    "Chronic conditions are seeded probabilistically from age, family history, smoking "
    "status, and BMI — producing realistic clusters of comorbidities."
)
add_code(
    "def comorbidity_seeds(age, fam_hx, smoker, bmi) -> set[str]:\n"
    "    # age >= 45: seeds HTN (30%), T2DM (20%), hyperlipidemia (35%)\n"
    "    # age >= 60: seeds COPD if smoker, hypothyroidism (25%), OA (40%)"
)

# ===========================================================================
# 6. CLI
# ===========================================================================
doc.add_heading("6.  Command-Line Interface", level=1)
doc.add_paragraph("All maintenance tasks are driven through cli.py.")
add_table(
    ["Command", "Purpose"],
    [
        ["generate --patients N --years Y", "Generate a fresh dataset."],
        ["stats", "Print dataset statistics."],
        ["export --format {json|fhir|text|all} --limit N --output-dir DIR", "Export records in one or all formats."],
        ["show --mrn MRN########", "Print a single patient's full chart."],
        ["list-conditions", "List all available condition codes."],
        ["add-spike --condition NAME --month M --n N", "Inject a seasonal disease spike."],
        ["advance --months M", "Advance time, adding follow-up visits for chronic patients."],
    ],
    widths=[3.4, 3.1],
)
doc.add_paragraph("Example — inject 300 extra influenza visits in January:")
add_code("python cli.py add-spike --condition influenza --month 1 --n 300")
doc.add_paragraph("Example — advance the timeline by 6 months:")
add_code("python cli.py advance --months 6")

# ===========================================================================
# 7. EXPORT FORMATS
# ===========================================================================
doc.add_heading("7.  Export Formats", level=1)
doc.add_paragraph("Records can be exported in three complementary formats.")

doc.add_heading("JSON (per patient)", level=3)
doc.add_paragraph(
    "A full denormalized bundle per patient — one file equals one patient's entire record, "
    "including chronic conditions and every visit with its vitals, diagnoses, prescriptions, "
    "and labs."
)

doc.add_heading("FHIR R4 Bundle", level=3)
doc.add_paragraph("A standard FHIR R4 Bundle per patient containing:")
add_bullets([
    "Patient resource.",
    "Encounter resource per visit.",
    "Observation resources for vitals (with LOINC codes).",
    "Condition resources for diagnoses (ICD-10).",
    "MedicationRequest resources for prescriptions.",
    "Observation resources for labs (with LOINC codes and reference ranges).",
])

doc.add_heading("Plain-Text Clinical Notes", level=3)
doc.add_paragraph("An LLM-ready chart summary — ideal as direct context for language models.")
add_code(
    "PATIENT CHART SUMMARY\n"
    "MRN: ... | Name: ... | Age: ... | Sex: ...\n"
    "FAMILY HISTORY: ...\n"
    "ACTIVE CHRONIC CONDITIONS: [ICD10] Description — Onset: date (Controlled/Uncontrolled)\n\n"
    "VISIT HISTORY (N total visits)\n"
    "DATE: 2022-03-06 [Follow Up] — Provider: Dr. James O'Brien, MD\n"
    "CHIEF COMPLAINT: Shortness of breath, worsening COPD\n"
    "VITALS: BP 129/88 | HR 74 | Temp 98.4°F | SpO2 93% | BMI 20.2 | Pain 1/10\n"
    "ASSESSMENT: J44.1 – COPD with acute exacerbation\n"
    "  Rx [Refill]: Albuterol inhaler 2 puffs Q4H PRN ×Ongoing\n"
    "  LABS: WBC 8.02 K/uL (Normal) | FEV1 40.13 %predicted ◄\n"
    "FOLLOW-UP: Return in 30 days"
)

# ===========================================================================
# 8. EXTENSIBLE SCHEMA
# ===========================================================================
doc.add_heading("8.  Extensible Schema Architecture", level=1)
doc.add_paragraph(
    "Beyond the core ORM, the project documents a JSON-driven, modular schema registry that "
    "lets teams extend the data model without editing core class bodies — and keeps Alembic "
    "autogenerate working cleanly."
)

doc.add_heading("Core concept", level=3)
doc.add_paragraph(
    "Declarative JSON is the source of truth — not Python class bodies. A registry merges "
    "schemas across modules, a ClassFactory builds the SQLAlchemy classes, and Alembic then "
    "sees ordinary metadata."
)

doc.add_heading("Modular layout", level=3)
add_bullets([
    "base_module — the core entities and relationships.",
    "ontology_module — adds ontology columns (e.g. SNOMED tags) and new entities.",
    "clinical_module — adds clinical extensions such as billing CPT codes or lab review fields.",
])
doc.add_paragraph(
    "Each module ships a manifest.json declaring its name, version, dependencies, and "
    "priority, plus a schema/ folder split into entities/ (columns + indexes) and "
    "relationships/ (relationship definitions)."
)

doc.add_heading("Four-phase load order", level=3)
doc.add_paragraph(
    "Entities and relationships live in separate files because columns only need their own "
    "table to exist, while relationships need both sides mapped. The registry loads in four "
    "phases, then the factory builds classes in two passes:"
)
add_table(
    ["Phase", "What happens"],
    [
        ["Phase 1 — Entity schemas", "Load all modules' columns; every tablename and column is now known."],
        ["Phase 2 — Merge entities", "Produce one merged column+index spec per entity (later module wins, logged)."],
        ["Phase 3 — Relationship schemas", "Validate all relationship targets against the complete entity set."],
        ["Phase 4 — Merge relationships", "Merge relationship specs (later module wins, logged)."],
        ["Factory Pass 1", "Create mapped classes (columns + indexes only)."],
        ["Factory Pass 2", "Wire relationships — no forward references or deferred resolution needed."],
    ],
    widths=[2.2, 4.3],
)

doc.add_heading("Merge collision rules", level=3)
add_table(
    ["Collision", "Resolution"],
    [
        ["Same column name, different modules", "Later module wins; warning logged."],
        ["Same relationship name, different modules", "Later module wins; warning logged."],
        ["Extension tries to rename tablename", "Hard error, blocked."],
        ["Relationship targets unknown entity", "Hard error at Phase 3 load time."],
        ["Circular module dependency", "Hard error from topological sort."],
    ],
    widths=[3.2, 3.3],
)

doc.add_heading("Bootstrap & Alembic", level=3)
doc.add_paragraph(
    "A single bootstrap() routine runs the full registry + factory sequence and must run in "
    "every process (CLI, Alembic, workers, API server) — the registry is a runtime construct, "
    "not persisted state. Adding a column is then a three-step workflow:"
)
add_code(
    "# 1. Add the column to the module's entities/*.json schema\n"
    "# 2. Autogenerate the migration\n"
    'alembic revision --autogenerate -m "clinical_module: add encounter_duration_min"\n'
    "# 3. Apply it\n"
    "alembic upgrade head"
)

# ===========================================================================
# 9. STATISTICS
# ===========================================================================
doc.add_heading("9.  Dataset Statistics", level=1)
doc.add_paragraph("The shipped 10,000-patient dataset contains:")
add_table(
    ["Metric", "Count"],
    [
        ["Patients", "10,000"],
        ["Visits", "165,972"],
        ["Diagnoses", "165,972"],
        ["Prescriptions", "170,267"],
        ["Lab Results", "777,868"],
        ["Avg visits / patient", "16.6"],
    ],
    widths=[2.6, 2.0],
)

doc.add_heading("Top Diagnoses", level=3)
add_table(
    ["ICD-10", "Description", "Count"],
    [
        ["I10", "Essential hypertension", "21,051"],
        ["E11.9", "Type 2 diabetes mellitus", "15,804"],
        ["E78.5", "Hyperlipidemia", "14,585"],
        ["Z00.00", "General adult medical examination", "14,187"],
        ["J06.9", "Acute upper respiratory infection", "11,860"],
        ["M19.90", "Osteoarthritis", "11,568"],
        ["J11.1", "Influenza", "9,056"],
        ["Z00.129", "Well-child visit", "7,762"],
        ["E03.9", "Hypothyroidism", "6,144"],
        ["J44.1", "COPD with acute exacerbation", "5,955"],
    ],
    widths=[1.1, 3.6, 1.2],
)

doc.add_heading("Age Distribution", level=3)
add_table(
    ["Age Band", "Patients", "Share"],
    [
        ["0–12", "2,059", "20.6%"],
        ["13–17", "814", "8.1%"],
        ["18–35", "1,689", "16.9%"],
        ["36–50", "1,415", "14.2%"],
        ["51–65", "1,425", "14.3%"],
        ["66+", "2,600", "26.0%"],
    ],
    widths=[1.8, 1.6, 1.4],
)

# ===========================================================================
# 10. ROADMAP
# ===========================================================================
doc.add_heading("10.  Roadmap & Possible Extensions", level=1)
add_bullets([
    "Agentic care program layer — a LangGraph agent with SQLite tools querying this dataset.",
    "Narrative generation — LLM-generated SOAP-note text per visit.",
    "Care-gap detection — flag patients overdue for preventive visits or with uncontrolled chronic conditions.",
    "Risk stratification — an ML model predicting hospitalization risk from visit patterns.",
    "FHIR server — wrap the exporters in a HAPI FHIR-compatible REST API.",
    "Ontology module — add SNOMED CT codes to diagnosis records.",
    "Billing module — add CPT codes, RVUs, and insurance-claim simulation.",
])

# ---- footer ---------------------------------------------------------------
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Family Medicine Synthetic Dataset — Feature Guide")
fr.font.size = Pt(8)
fr.font.color.rgb = GREY

doc.save(r"C:\Users\arsal\projects\hdh\Feature_Guide.docx")
print("Saved Feature_Guide.docx")
